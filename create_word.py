from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ========== 设置默认样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ========== 页面设置 ==========
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ========== 辅助函数 ==========
def add_title(text, level=0):
    """添加标题"""
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x21, 0x25, 0x29)
    else:
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x43, 0x61, 0xEE)

def add_para(text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.bold = bold
    return p

def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'

def add_code_block(code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    # 添加灰色背景效果
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_separator():
    """添加分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

# ========== 文档标题 ==========
add_title('个人博客网站毕业设计')
add_title('答辩问题回答参考', level=1)

# 基本信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('答辩人：周驰    完成日期：2025年12月')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

doc.add_page_break()

# ========== 问题1 ==========
add_title('问题一：你的博客用到了哪些技术？主要功能有哪些？', level=2)

add_para('一、技术栈', bold=True)

add_para('1. 前端技术', bold=True)
add_bullet('HTML5 / CSS3：构建语义化标签结构和响应式布局，使用 CSS Flexbox 和 Grid 实现现代化页面布局')
add_bullet('JavaScript（ES6+）：实现交互功能和动态内容加载，使用 Fetch API 进行前后端异步通信')
add_bullet('marked.js 库：解析 Markdown 格式的文章内容，支持实时预览和文章目录生成')

add_para('2. 后端技术', bold=True)
add_bullet('Node.js：JavaScript 运行时环境，提供异步非阻塞 I/O 处理能力')
add_bullet('Express.js：轻量级 Web 应用框架，提供路由管理、中间件机制等核心功能')
add_bullet('JWT（JSON Web Token）：实现无状态用户认证，令牌有效期 24 小时')
add_bullet('Multer：文件上传处理中间件，支持文件类型验证和大小限制')
add_bullet('bcryptjs：密码加密存储，使用 10 轮盐值哈希算法防止彩虹表攻击')

add_para('3. 数据库技术', bold=True)
add_bullet('SQLite：嵌入式关系型数据库，无需独立服务器，零配置易于部署')
add_bullet('数据库包含 4 个核心表：users（用户）、articles（文章）、comments（评论）、gallery_items（画廊作品）')

add_separator()

add_para('二、主要功能', bold=True)

add_para('1. 用户系统', bold=True)
add_bullet('用户注册：支持用户名、密码、邮箱注册')
add_bullet('用户登录：验证身份后返回 JWT 令牌，前端存储在 localStorage')
add_bullet('安全退出：清除本地存储的令牌信息')
add_bullet('登录状态保持：页面刷新后自动恢复登录状态')

add_para('2. 文章管理', bold=True)
add_bullet('发布文章：支持 Markdown 语法编辑，提供工具栏快捷插入')
add_bullet('编辑文章：作者可修改已发布的文章')
add_bullet('删除文章：作者可删除自己的文章，同时删除关联评论')
add_bullet('草稿功能：文章可保存为草稿，不公开显示')
add_bullet('分类管理：文章可按分类浏览和筛选')
add_bullet('标签系统：支持多标签标记文章')
add_bullet('搜索功能：按标题、内容、标签搜索文章')
add_bullet('分页展示：支持分页浏览文章列表')

add_para('3. 图片上传', bold=True)
add_bullet('支持 JPG、PNG、GIF 格式图片上传，最大 2MB')
add_bullet('支持拖拽上传和点击上传两种方式')
add_bullet('上传前可预览图片')
add_bullet('上传后可插入到文章内容中')
add_bullet('上传图片可设为文章封面')

add_para('4. 评论系统', bold=True)
add_bullet('匿名评论：无需登录即可发表评论')
add_bullet('评论列表：按时间倒序展示')
add_bullet('实时更新：发表评论后自动刷新显示')

add_para('5. 摄影画廊', bold=True)
add_bullet('摄影作品分类展示（风光、人像、街头、自然、微距）')
add_bullet('支持按分类筛选作品')
add_bullet('作品详情查看，支持下载图片')

add_para('6. 其他功能', bold=True)
add_bullet('夜间模式：一键切换，本地保存偏好设置')
add_bullet('响应式设计：适配桌面端、平板和手机屏幕')
add_bullet('回到顶部：长页面快捷返回顶部')
add_bullet('文章目录：自动生成文章标题导航')

doc.add_page_break()

# ========== 问题2 ==========
add_title('问题二：文章能不能修改、删除、存草稿？怎么实现？', level=2)

add_para('可以。文章支持完整的增删改查（CRUD）操作和草稿功能。', bold=True)

add_separator()

add_para('一、修改文章', bold=True)

add_para('前端实现：', bold=True)
add_bullet('文章卡片上显示"编辑"按钮（仅文章作者可见）')
add_bullet('点击编辑按钮后，通过 GET 请求获取文章完整数据')
add_bullet('将文章标题、内容、分类、标签等数据填入编辑表单')
add_bullet('修改完成后点击"保存修改"，发送 PUT 请求更新文章')

add_para('后端实现：', bold=True)
add_bullet('路由：PUT /api/articles/:id')
add_bullet('首先验证 JWT 令牌，确认用户已登录')
add_bullet('检查文章是否存在，以及当前用户是否为文章作者')
add_bullet('执行 UPDATE SQL 语句更新数据库记录')
add_bullet('自动更新 updated_at 字段为当前时间')

add_para('核心代码：', bold=True)
code = '''// 更新文章（需要认证）
app.put('/api/articles/:id', authenticateToken, (req, res) => {
    const articleId = req.params.id;
    const { title, content, category, tags, status, cover_image } = req.body;
    const userId = req.user.userId;

    // 检查文章是否存在
    db.get('SELECT * FROM articles WHERE id = ?', [articleId], (err, article) => {
        if (!article) return res.status(404).json({ error: '文章不存在' });
        // 检查权限：只有作者可以编辑
        if (article.author_id !== userId) 
            return res.status(403).json({ error: '无权编辑此文章' });
        // 更新文章
        db.run('UPDATE articles SET title=?, content=?, category=?, tags=?, '
            + 'status=?, cover_image=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
            [title, content, category, tags, status, cover_image, articleId]);
    });
});'''
add_code_block(code)

add_separator()

add_para('二、删除文章', bold=True)

add_para('前端实现：', bold=True)
add_bullet('文章卡片上显示"删除"按钮（仅文章作者可见）')
add_bullet('点击删除按钮后，弹出确认对话框："确定要删除这篇文章吗？此操作不可撤销。"')
add_bullet('用户确认后，发送 DELETE 请求到后端')

add_para('后端实现：', bold=True)
add_bullet('路由：DELETE /api/articles/:id')
add_bullet('验证 JWT 令牌和文章作者权限')
add_bullet('使用数据库事务（BEGIN TRANSACTION）确保数据一致性')
add_bullet('先删除文章的所有关联评论，再删除文章本身')
add_bullet('如果任一操作失败，执行 ROLLBACK 回滚事务')

add_para('核心代码：', bold=True)
code = '''// 删除文章（需要认证）
app.delete('/api/articles/:id', authenticateToken, (req, res) => {
    // ...验证权限...
    db.serialize(() => {
        db.run('BEGIN TRANSACTION');
        // 先删除评论
        db.run('DELETE FROM comments WHERE article_id = ?', [articleId]);
        // 再删除文章
        db.run('DELETE FROM articles WHERE id = ?', [articleId]);
        db.run('COMMIT');
    });
});'''
add_code_block(code)

add_separator()

add_para('三、存草稿', bold=True)

add_para('实现方式：', bold=True)
add_bullet('articles 表中有一个 status 字段，取值可以是 "published"（已发布）或 "draft"（草稿）')
add_bullet('创建文章时，如果传入 status = "draft"，则文章保存为草稿状态')
add_bullet('文章列表查询时，SQL 语句中带有 WHERE a.status = "published" 条件，草稿不会出现在公开列表中')
add_bullet('用户可以在个人中心查看和管理自己的所有文章，包括草稿')
add_bullet('草稿可以随时编辑并发布')

doc.add_page_break()

# ========== 问题3 ==========
add_title('问题三：文章详情页会显示阅读量、发布时间吗？', level=2)

add_para('会显示。文章详情页展示的信息非常完整，包括以下内容：', bold=True)

add_separator()

add_para('一、显示的信息', bold=True)

add_para('1. 阅读量', bold=True)
add_bullet('每次用户访问文章详情时，后端自动执行 UPDATE articles SET views = views + 1 WHERE id = ?')
add_bullet('阅读量实时累加，刷新页面会再次增加')
add_bullet('在文章卡片和详情页都会显示阅读量图标和数字')

add_para('2. 发布时间', bold=True)
add_bullet('显示 created_at 字段，格式化为中文日期时间格式')
add_bullet('例如："2025年12月10日 14:30"')

add_para('3. 更新时间', bold=True)
add_bullet('文章编辑后 updated_at 字段自动更新为当前时间')
add_bullet('数据库中使用 DEFAULT CURRENT_TIMESTAMP 自动记录时间')

add_para('4. 其他信息', bold=True)
add_bullet('作者：通过 JOIN 查询关联 users 表获取用户名')
add_bullet('分类：显示文章所属分类')
add_bullet('标签：显示文章的标签信息')
add_bullet('点赞数：显示 likes 字段，用户可点击点赞按钮')
add_bullet('评论数：显示评论数量及完整的评论列表')
add_bullet('封面图：如果文章设置了封面图，会显示在详情页顶部')

add_separator()

add_para('二、后端实现代码', bold=True)

code = '''// 获取单篇文章
app.get('/api/articles/:id', (req, res) => {
    const articleId = req.params.id;

    // 增加阅读量
    db.run('UPDATE articles SET views = views + 1 WHERE id = ?', [articleId]);

    // 查询文章详情（关联作者信息）
    db.get('SELECT a.*, u.username as author_name FROM articles a 
            LEFT JOIN users u ON a.author_id = u.id WHERE a.id = ?', 
        [articleId], (err, article) => {
        // 同时查询评论列表
        db.all('SELECT * FROM comments WHERE article_id = ? 
                ORDER BY created_at DESC', [articleId], (err, comments) => {
            res.json({ ...article, comments });
        });
    });
});'''
add_code_block(code)

add_separator()

add_para('三、前端展示代码', bold=True)

code = '''// 前端渲染文章详情
function displayArticleDetail(article) {
    articleContent.innerHTML = `
        <h1 class="article-detail-title">${article.title}</h1>
        <div class="article-detail-meta">
            <span>👤 ${article.author_name}</span>
            <span>📂 ${article.category}</span>
            <span>📅 ${formatDate(article.created_at)}</span>
            <span>👁️ ${article.views}</span>
            <span>❤️ ${article.likes || 0}</span>
            ${article.tags ? `<span>🏷️ ${article.tags}</span>` : ''}
        </div>
        ${article.cover_image ? `<img src="..." class="article-detail-cover">` : ''}
        <div class="article-detail-content">${htmlContent}</div>
        <!-- 评论区域 -->
    `;
}'''
add_code_block(code)

doc.add_page_break()

# ========== 问题4 ==========
add_title('问题四：前后端是怎么传数据的？', level=2)

add_para('采用 RESTful API + JSON 格式进行前后端通信。', bold=True)

add_separator()

add_para('一、通信方式', bold=True)

add_para('前端使用浏览器内置的 Fetch API 发送 HTTP 请求，后端使用 Express.js 接收请求并返回 JSON 格式的响应数据。这是一种典型的 AJAX（Asynchronous JavaScript and XML）异步通信模式。', indent=True)

add_separator()

add_para('二、通信流程图', bold=True)

code = '''前端（浏览器）                   后端（Node.js服务器）
    │                                  │
    │  ── HTTP 请求（Fetch API）──▶    │
    │     GET    /api/articles          │  获取文章列表
    │     GET    /api/articles/:id      │  获取文章详情
    │     POST   /api/login             │  用户登录
    │     POST   /api/register          │  用户注册
    │     POST   /api/articles          │  创建文章
    │     PUT    /api/articles/:id      │  更新文章
    │     DELETE /api/articles/:id      │  删除文章
    │     POST   /api/articles/:id/like │  点赞文章
    │     POST   /api/upload/image      │  上传图片
    │                                  │
    │  ◀── JSON 响应 ──────────────    │
    │     { articles: [...],            │
    │       pagination: {...} }         │
    │     { message: "登录成功",        │
    │       token: "xxx",               │
    │       user: {...} }               │
    │     { message: "点赞成功" }       │'''
add_code_block(code)

add_separator()

add_para('三、前端发送请求示例', bold=True)

add_para('1. 获取文章列表（GET 请求）：', bold=True)
code = '''async function loadHomePage() {
    const response = await fetch(`${API_BASE_URL}/articles?page=1&limit=10`);
    const data = await response.json();
    // data.articles  - 文章数组
    // data.pagination - 分页信息
}'''
add_code_block(code)

add_para('2. 用户登录（POST 请求，发送 JSON 数据）：', bold=True)
code = '''async function handleLogin(event) {
    const response = await fetch(`${API_BASE_URL}/login`, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({
            username: 'admin',
            password: 'admin123'
        })
    });
    const data = await response.json();
    // data.token - JWT令牌
    // data.user  - 用户信息
}'''
add_code_block(code)

add_para('3. 创建文章（POST 请求，需要认证）：', bold=True)
code = '''async function handleWriteArticle(event) {
    const token = localStorage.getItem('token');
    const response = await fetch(`${API_BASE_URL}/articles`, {
        method: 'POST',
        headers: {
            'Content-Type': 'application/json',
            'Authorization': `Bearer ${token}`  // JWT认证
        },
        body: JSON.stringify({
            title: '文章标题',
            content: '文章内容（支持Markdown）',
            category: '技术',
            tags: 'Node.js,JavaScript',
            cover_image: 'https://...'
        })
    });
}'''
add_code_block(code)

add_separator()

add_para('四、后端处理请求示例', bold=True)

code = '''// 后端路由处理
app.get('/api/articles', (req, res) => {
    // 1. 解析请求参数
    const page = parseInt(req.query.page) || 1;
    const limit = parseInt(req.query.limit) || 10;
    
    // 2. 查询数据库
    db.all('SELECT * FROM articles ... LIMIT ? OFFSET ?', 
        [limit, offset], (err, articles) => {
        
        // 3. 返回JSON响应
        res.json({
            articles: articles,
            pagination: {
                page: page,
                total: countResult.total,
                totalPages: Math.ceil(countResult.total / limit)
            }
        });
    });
});'''
add_code_block(code)

add_separator()

add_para('五、数据格式说明', bold=True)

add_para('请求格式：', bold=True)
add_bullet('GET 请求：参数通过 URL 查询字符串传递（?page=1&limit=10）')
add_bullet('POST/PUT 请求：数据以 JSON 格式放在请求体（body）中')
add_bullet('需要认证的请求：在 HTTP 头部添加 Authorization: Bearer <token>')
add_bullet('文件上传：使用 FormData 格式，由 multer 中间件解析')

add_para('响应格式：', bold=True)
add_bullet('所有响应均为 JSON 格式')
add_bullet('成功响应：返回请求的数据或成功消息')
add_bullet('错误响应：返回 { error: "错误描述" }')
add_bullet('HTTP 状态码：200（成功）、400（参数错误）、401（未登录）、403（无权限）、404（不存在）、500（服务器错误）')

add_separator()

add_para('六、API 接口汇总', bold=True)

add_para('用户认证：', bold=True)
add_bullet('POST /api/register — 用户注册')
add_bullet('POST /api/login — 用户登录')

add_para('文章管理：', bold=True)
add_bullet('GET /api/articles — 获取文章列表（分页）')
add_bullet('GET /api/articles/:id — 获取单篇文章详情')
add_bullet('POST /api/articles — 创建新文章（需认证）')
add_bullet('PUT /api/articles/:id — 更新文章（需认证）')
add_bullet('DELETE /api/articles/:id — 删除文章（需认证）')
add_bullet('GET /api/articles/search/:query — 搜索文章')
add_bullet('POST /api/articles/:id/like — 点赞文章')

add_para('分类与评论：', bold=True)
add_bullet('GET /api/categories — 获取文章分类')
add_bullet('GET /api/articles/category/:category — 按分类获取文章')
add_bullet('POST /api/articles/:id/comments — 添加评论')

add_para('图片与画廊：', bold=True)
add_bullet('POST /api/upload/image — 上传图片（需认证）')
add_bullet('GET /api/gallery/items — 获取画廊作品列表')
add_bullet('POST /api/gallery/items — 上传画廊作品（需认证）')

add_para('用户中心：', bold=True)
add_bullet('GET /api/user/profile — 获取用户信息（需认证）')
add_bullet('GET /api/user/articles — 获取用户文章列表（需认证）')

# ========== 保存文件 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "答辩问题回答参考.docx")
doc.save(output_path)
print(f"Word文档已生成: {output_path}")
