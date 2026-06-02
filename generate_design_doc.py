# -*- coding: utf-8 -*-
"""
生成博客网站设计说明书
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def add_para(doc, text, bold=False, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # 数据行
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    
    return table


def create_design_doc():
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()
    
    add_para(doc, "个人摄影博客网站", bold=True, font_size=28, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "设计说明书", bold=True, font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_para(doc, "项目名称：个人摄影博客网站", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "开发技术：Node.js + Express + SQLite + 原生前端", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "开发人员：周驰", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "班　　级：计网231", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "学　　号：2023062101", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "指导老师：李志敏", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "完成日期：2026年5月", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()
    
    # ========== 目录页 ==========
    add_para(doc, "目  录", bold=True, font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    toc_items = [
        "一、项目概述",
        "二、需求分析",
        "三、系统架构设计",
        "四、功能模块设计",
        "五、数据库设计",
        "六、接口设计",
        "七、前端界面设计",
        "八、安全设计",
        "九、部署与运行",
        "十、总结与展望",
    ]
    for item in toc_items:
        add_para(doc, item, font_size=14, space_after=8)
    
    doc.add_page_break()
    
    # ========== 一、项目概述 ==========
    add_heading(doc, "一、项目概述", level=1)
    
    add_heading(doc, "1.1 项目背景", level=2)
    add_para(doc, "随着互联网技术的快速发展和社交媒体的普及，个人博客作为一种展示自我、分享知识的平台，依然具有独特的价值。"
             "摄影作为一种视觉艺术形式，需要一个能够充分展示图片和文字内容的平台。"
             "本系统旨在开发一个功能完善、界面美观的个人摄影博客网站，为摄影爱好者提供一个展示作品、分享经验、交流互动的平台。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "1.2 项目目标", level=2)
    add_para(doc, "本项目的目标是构建一个基于Node.js的个人摄影博客网站，实现以下核心目标：", font_size=12, first_line_indent=0.74)
    goals = [
        "提供一个美观、响应式的博客展示平台，支持文章浏览、搜索和分类查看",
        "实现完整的用户注册、登录和权限管理系统",
        "支持文章的创建、编辑、删除等完整CRUD操作，并支持Markdown编辑",
        "提供评论互动功能，增强用户参与度",
        "集成画廊作品集功能，方便展示摄影作品",
        "支持夜间模式切换，提升用户体验",
        "具备良好的安全性和可扩展性",
    ]
    for g in goals:
        add_para(doc, f"• {g}", font_size=12, space_after=3)
    
    add_heading(doc, "1.3 开发环境", level=2)
    
    env_data = [
        ["操作系统", "Windows 11"],
        ["运行环境", "Node.js 18+"],
        ["开发工具", "Visual Studio Code"],
        ["数据库", "SQLite 3"],
        ["版本控制", "Git"],
        ["测试浏览器", "Chrome、Edge、Firefox"],
    ]
    add_table(doc, ["项目", "说明"], env_data, [4, 10])
    
    doc.add_page_break()
    
    # ========== 二、需求分析 ==========
    add_heading(doc, "二、需求分析", level=1)
    
    add_heading(doc, "2.1 功能需求", level=2)
    
    add_para(doc, "（1）用户管理模块", bold=True, font_size=12)
    add_para(doc, "• 用户注册：支持新用户通过用户名、邮箱、密码进行注册\n"
             "• 用户登录：支持已注册用户通过用户名和密码登录\n"
             "• 身份认证：采用JWT（JSON Web Token）进行用户身份认证\n"
             "• 个人中心：展示用户信息、文章统计数据和文章列表", font_size=12, first_line_indent=0.74)
    
    add_para(doc, "（2）文章管理模块", bold=True, font_size=12)
    add_para(doc, "• 文章列表：分页展示已发布的文章，支持按时间排序\n"
             "• 文章详情：展示文章完整内容，支持Markdown渲染\n"
             "• 文章创建：支持Markdown编辑器，提供工具栏和实时预览\n"
             "• 文章编辑：支持作者对自己文章的编辑修改\n"
             "• 文章删除：支持作者对自己文章的删除操作\n"
             "• 文章分类：支持对文章进行分类管理\n"
             "• 文章标签：支持为文章添加多个标签\n"
             "• 封面图片：支持为文章设置封面图", font_size=12, first_line_indent=0.74)
    
    add_para(doc, "（3）评论互动模块", bold=True, font_size=12)
    add_para(doc, "• 评论发布：支持访客对文章发表评论\n"
             "• 评论展示：按时间倒序展示文章评论\n"
             "• 点赞功能：支持对文章进行点赞", font_size=12, first_line_indent=0.74)
    
    add_para(doc, "（4）搜索与分类模块", bold=True, font_size=12)
    add_para(doc, "• 文章搜索：支持按标题、内容、标签进行全文搜索\n"
             "• 分类浏览：按分类展示文章列表\n"
             "• 热门文章：展示阅读量最高的文章", font_size=12, first_line_indent=0.74)
    
    add_para(doc, "（5）画廊模块", bold=True, font_size=12)
    add_para(doc, "• 作品展示：以网格形式展示摄影作品\n"
             "• 作品分类：支持按类别筛选作品\n"
             "• 作品详情：查看作品大图和详细信息\n"
             "• 作品上传：支持登录用户上传摄影作品", font_size=12, first_line_indent=0.74)
    
    add_para(doc, "（6）系统功能", bold=True, font_size=12)
    add_para(doc, "• 夜间模式：支持一键切换日间/夜间主题\n"
             "• 响应式设计：适配桌面端和移动端\n"
             "• 图片上传：支持上传图片并插入文章或设为封面\n"
             "• 回到顶部：长页面快速回到顶部", font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "2.2 非功能需求", level=2)
    nf_reqs = [
        "性能要求：页面加载时间不超过3秒，API响应时间不超过500ms",
        "安全要求：用户密码加密存储，API接口需认证访问，防止SQL注入",
        "可用性要求：系统7×24小时稳定运行，界面友好易用",
        "兼容性要求：支持主流浏览器最新版本",
        "可维护性要求：代码结构清晰，注释完整，便于后续维护和扩展",
    ]
    for r in nf_reqs:
        add_para(doc, f"• {r}", font_size=12, space_after=3)
    
    doc.add_page_break()
    
    # ========== 三、系统架构设计 ==========
    add_heading(doc, "三、系统架构设计", level=1)
    
    add_heading(doc, "3.1 系统架构图", level=2)
    add_para(doc, "本系统采用前后端分离的B/S（Browser/Server）架构，前端负责页面展示和用户交互，"
             "后端负责业务逻辑处理和数据处理。整体架构分为三层：", font_size=12, first_line_indent=0.74)
    
    arch_data = [
        ["表现层（前端）", "HTML5 + CSS3 + JavaScript", "页面展示、用户交互、路由控制"],
        ["业务逻辑层（后端）", "Node.js + Express", "API接口、业务处理、认证授权"],
        ["数据访问层（数据库）", "SQLite 3", "数据存储、数据查询、数据管理"],
    ]
    add_table(doc, ["层次", "技术", "职责"], arch_data, [3.5, 5, 5.5])
    
    doc.add_paragraph()
    add_para(doc, "前端通过HTTP请求调用后端RESTful API，后端处理请求后返回JSON格式数据。"
             "前端使用原生JavaScript进行DOM操作和异步数据加载，实现单页应用（SPA）效果。", font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "3.2 技术选型", level=2)
    
    tech_data = [
        ["Node.js", "JavaScript运行时", "提供服务器端运行环境，事件驱动、非阻塞I/O"],
        ["Express", "Web应用框架", "简化路由、中间件、请求处理"],
        ["SQLite 3", "嵌入式数据库", "轻量级、无需单独部署、适合小型应用"],
        ["JWT", "身份认证", "无状态认证、安全可靠"],
        ["bcryptjs", "密码加密", "密码哈希存储、防止泄露"],
        ["multer", "文件上传", "处理图片上传、文件存储"],
        ["marked", "Markdown解析", "将Markdown文本渲染为HTML"],
        ["Font Awesome", "图标库", "提供丰富的UI图标"],
    ]
    add_table(doc, ["技术/工具", "类型", "作用"], tech_data, [3, 4, 7])
    
    doc.add_page_break()
    
    # ========== 四、功能模块设计 ==========
    add_heading(doc, "四、功能模块设计", level=1)
    
    add_heading(doc, "4.1 模块划分", level=2)
    add_para(doc, "系统共分为以下六大功能模块：", font_size=12)
    
    modules = [
        ["用户管理模块", "注册、登录、认证、个人中心"],
        ["文章管理模块", "文章CRUD、Markdown编辑、分类标签、封面图"],
        ["评论互动模块", "评论发布与展示、点赞"],
        ["搜索分类模块", "全文搜索、分类浏览、热门推荐"],
        ["画廊模块", "作品展示、上传、分类筛选"],
        ["系统功能模块", "夜间模式、响应式设计、图片上传"],
    ]
    add_table(doc, ["模块名称", "子功能"], modules, [4, 10])
    
    add_heading(doc, "4.2 用户管理模块", level=2)
    add_para(doc, "用户管理模块负责用户的注册、登录和身份认证。用户注册时需提供用户名和密码，"
             "密码经过bcrypt加密后存储到数据库。登录成功后，服务器生成JWT令牌返回给客户端，"
             "客户端将令牌存储在localStorage中，后续请求在Authorization头中携带令牌进行身份验证。"
             "个人中心页面展示用户的基本信息、文章统计数据（文章数、总阅读量、总评论数）以及用户发布的文章列表。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "4.3 文章管理模块", level=2)
    add_para(doc, "文章管理模块是系统的核心功能。支持文章的创建、读取、更新和删除（CRUD）操作。"
             "文章编辑器集成了Markdown语法支持，提供加粗、斜体、标题、列表、引用、代码块等工具栏按钮，"
             "并支持实时预览功能。文章可以设置分类和标签，方便管理和检索。"
             "文章列表支持分页展示，每页显示10篇文章。文章详情页展示完整内容，并自动生成目录导航。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "4.4 评论互动模块", level=2)
    add_para(doc, "评论互动模块允许访客对文章发表评论，评论需要填写昵称和评论内容。"
             "评论按发布时间倒序排列，最新评论显示在最前面。"
             "点赞功能允许用户对喜欢的文章进行点赞，点赞数实时更新。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "4.5 搜索分类模块", level=2)
    add_para(doc, "搜索功能支持按文章标题、内容和标签进行模糊匹配搜索，搜索结果支持分页展示。"
             "分类功能从数据库中提取所有已发布文章的分类信息，以卡片形式展示各分类及其文章数量，"
             "点击分类可查看该分类下的所有文章。侧边栏展示热门文章列表，按阅读量排序。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "4.6 画廊模块", level=2)
    add_para(doc, "画廊模块以网格布局展示摄影作品，支持按类别筛选。"
             "登录用户可以上传自己的摄影作品，上传时需要提供标题、描述和分类信息。"
             "作品详情页展示大图、描述信息和浏览数据。", 
             font_size=12, first_line_indent=0.74)
    
    doc.add_page_break()
    
    # ========== 五、数据库设计 ==========
    add_heading(doc, "五、数据库设计", level=1)
    
    add_heading(doc, "5.1 数据库选型", level=2)
    add_para(doc, "本系统采用SQLite 3作为数据库。SQLite是一个轻量级的嵌入式关系数据库引擎，"
             "不需要单独的服务器进程，数据存储在一个文件中，非常适合小型Web应用和个人项目。"
             "数据库文件为blog.db，位于项目根目录。", font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "5.2 数据表结构", level=2)
    
    add_para(doc, "（1）用户表（users）", bold=True, font_size=12)
    user_cols = ["id", "INTEGER", "主键，自增", "PRIMARY KEY AUTOINCREMENT"]
    user_cols2 = ["username", "TEXT", "用户名，唯一", "UNIQUE NOT NULL"]
    user_cols3 = ["password", "TEXT", "密码（bcrypt加密）", "NOT NULL"]
    user_cols4 = ["email", "TEXT", "邮箱，唯一", "UNIQUE"]
    user_cols5 = ["created_at", "DATETIME", "注册时间", "DEFAULT CURRENT_TIMESTAMP"]
    add_table(doc, ["字段名", "类型", "说明", "约束"], 
              [user_cols, user_cols2, user_cols3, user_cols4, user_cols5], [3, 3, 5, 5])
    
    doc.add_paragraph()
    add_para(doc, "（2）文章表（articles）", bold=True, font_size=12)
    art_cols = [
        ["id", "INTEGER", "主键，自增", "PRIMARY KEY AUTOINCREMENT"],
        ["title", "TEXT", "文章标题", "NOT NULL"],
        ["content", "TEXT", "文章内容（Markdown）", "NOT NULL"],
        ["author_id", "INTEGER", "作者ID", "NOT NULL, FOREIGN KEY"],
        ["category", "TEXT", "文章分类", "DEFAULT '未分类'"],
        ["tags", "TEXT", "标签（逗号分隔）", ""],
        ["cover_image", "TEXT", "封面图URL", "DEFAULT ''"],
        ["likes", "INTEGER", "点赞数", "DEFAULT 0"],
        ["views", "INTEGER", "阅读量", "DEFAULT 0"],
        ["status", "TEXT", "状态（published/draft）", "DEFAULT 'published'"],
        ["created_at", "DATETIME", "创建时间", "DEFAULT CURRENT_TIMESTAMP"],
        ["updated_at", "DATETIME", "更新时间", "DEFAULT CURRENT_TIMESTAMP"],
    ]
    add_table(doc, ["字段名", "类型", "说明", "约束"], art_cols, [3, 3, 5, 5])
    
    doc.add_paragraph()
    add_para(doc, "（3）评论表（comments）", bold=True, font_size=12)
    com_cols = [
        ["id", "INTEGER", "主键，自增", "PRIMARY KEY AUTOINCREMENT"],
        ["article_id", "INTEGER", "所属文章ID", "NOT NULL, FOREIGN KEY"],
        ["user_name", "TEXT", "评论者昵称", "NOT NULL"],
        ["content", "TEXT", "评论内容", "NOT NULL"],
        ["created_at", "DATETIME", "评论时间", "DEFAULT CURRENT_TIMESTAMP"],
    ]
    add_table(doc, ["字段名", "类型", "说明", "约束"], com_cols, [3, 3, 5, 5])
    
    doc.add_paragraph()
    add_para(doc, "（4）画廊作品表（gallery_items）", bold=True, font_size=12)
    gal_cols = [
        ["id", "INTEGER", "主键，自增", "PRIMARY KEY AUTOINCREMENT"],
        ["title", "TEXT", "作品标题", "NOT NULL"],
        ["description", "TEXT", "作品描述", ""],
        ["image_url", "TEXT", "图片URL", "NOT NULL"],
        ["category", "TEXT", "作品分类", "DEFAULT 'other'"],
        ["tags", "TEXT", "标签", ""],
        ["author_id", "INTEGER", "作者ID", "NOT NULL, FOREIGN KEY"],
        ["views", "INTEGER", "浏览量", "DEFAULT 0"],
        ["created_at", "DATETIME", "上传时间", "DEFAULT CURRENT_TIMESTAMP"],
    ]
    add_table(doc, ["字段名", "类型", "说明", "约束"], gal_cols, [3, 3, 5, 5])
    
    add_heading(doc, "5.3 实体关系（E-R）", level=2)
    add_para(doc, "系统中主要包含四个实体：用户（User）、文章（Article）、评论（Comment）和画廊作品（GalleryItem）。"
             "它们之间的关系如下：", font_size=12, first_line_indent=0.74)
    relations = [
        "用户与文章：一对多关系。一个用户可以发布多篇文章，一篇文章属于一个用户。",
        "文章与评论：一对多关系。一篇文章可以有多条评论，一条评论属于一篇文章。",
        "用户与画廊作品：一对多关系。一个用户可以上传多个作品，一个作品属于一个用户。",
    ]
    for r in relations:
        add_para(doc, f"• {r}", font_size=12, space_after=3)
    
    doc.add_page_break()
    
    # ========== 六、接口设计 ==========
    add_heading(doc, "六、接口设计", level=1)
    
    add_heading(doc, "6.1 接口规范", level=2)
    add_para(doc, "本系统采用RESTful API设计风格，所有接口返回JSON格式数据。"
             "接口地址以 /api 为前缀。需要认证的接口在请求头中携带 Authorization: Bearer <token>。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "6.2 接口列表", level=2)
    
    add_para(doc, "（1）用户相关接口", bold=True, font_size=12)
    api_users = [
        ["POST", "/api/register", "用户注册", "否"],
        ["POST", "/api/login", "用户登录", "否"],
        ["GET", "/api/user/profile", "获取个人信息", "是"],
        ["GET", "/api/user/articles", "获取我的文章列表", "是"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_users, [2, 5, 4, 2])
    
    doc.add_paragraph()
    add_para(doc, "（2）文章相关接口", bold=True, font_size=12)
    api_articles = [
        ["GET", "/api/articles", "获取文章列表（分页）", "否"],
        ["GET", "/api/articles/:id", "获取文章详情", "否"],
        ["POST", "/api/articles", "创建文章", "是"],
        ["PUT", "/api/articles/:id", "更新文章", "是"],
        ["DELETE", "/api/articles/:id", "删除文章", "是"],
        ["POST", "/api/articles/:id/like", "点赞文章", "否"],
        ["GET", "/api/articles/category/:category", "按分类获取文章", "否"],
        ["GET", "/api/articles/search/:query", "搜索文章", "否"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_articles, [2, 5, 4, 2])
    
    doc.add_paragraph()
    add_para(doc, "（3）评论相关接口", bold=True, font_size=12)
    api_comments = [
        ["POST", "/api/articles/:id/comments", "添加评论", "否"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_comments, [2, 5, 4, 2])
    
    doc.add_paragraph()
    add_para(doc, "（4）分类相关接口", bold=True, font_size=12)
    api_cats = [
        ["GET", "/api/categories", "获取所有分类", "否"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_cats, [2, 5, 4, 2])
    
    doc.add_paragraph()
    add_para(doc, "（5）画廊相关接口", bold=True, font_size=12)
    api_gallery = [
        ["GET", "/api/gallery/items", "获取作品列表（分页）", "否"],
        ["GET", "/api/gallery/items/:id", "获取作品详情", "否"],
        ["POST", "/api/gallery/items", "上传作品", "是"],
        ["DELETE", "/api/gallery/items/:id", "删除作品", "是"],
        ["GET", "/api/user/gallery-items", "获取我的作品列表", "是"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_gallery, [2, 5, 4, 2])
    
    doc.add_paragraph()
    add_para(doc, "（6）文件上传接口", bold=True, font_size=12)
    api_upload = [
        ["POST", "/api/upload/image", "上传图片", "是"],
    ]
    add_table(doc, ["方法", "路径", "说明", "需认证"], api_upload, [2, 5, 4, 2])
    
    doc.add_page_break()
    
    # ========== 七、前端界面设计 ==========
    add_heading(doc, "七、前端界面设计", level=1)
    
    add_heading(doc, "7.1 设计风格", level=2)
    add_para(doc, "前端界面采用现代化设计风格，以摄影艺术为主题。主色调为蓝色系（#4361ee），"
             "搭配紫色和粉色作为辅助色。页面布局清晰，视觉层次分明，注重图片展示效果。"
             "支持日间和夜间两种主题模式，用户可一键切换。", font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "7.2 页面结构", level=2)
    
    pages = [
        ["首页", "展示英雄区域、最新文章列表、侧边栏（分类、热门文章、关于我）"],
        ["文章详情页", "展示文章完整内容、目录导航、评论区"],
        ["分类页", "以卡片网格展示所有文章分类及文章数量"],
        ["登录页", "用户登录表单"],
        ["注册页", "用户注册表单"],
        ["写文章页", "Markdown编辑器、工具栏、图片上传、预览功能"],
        ["编辑文章页", "与写文章页类似，预填已有文章内容"],
        ["个人中心页", "用户信息、统计数据、我的文章列表"],
        ["画廊页", "作品网格展示、分类筛选、作品详情"],
    ]
    add_table(doc, ["页面", "功能说明"], pages, [3, 11])
    
    add_heading(doc, "7.3 响应式设计", level=2)
    add_para(doc, "页面采用CSS Flexbox和Grid布局实现响应式设计。在桌面端（≥1024px）显示完整布局，"
             "在平板端（768px-1023px）调整侧边栏和网格列数，在移动端（<768px）采用单列布局，"
             "导航栏折叠为汉堡菜单。确保在不同设备上都能获得良好的浏览体验。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "7.4 交互设计", level=2)
    interactions = [
        "页面切换：通过JavaScript控制页面显示/隐藏，实现单页应用效果，无需整页刷新",
        "实时预览：Markdown编辑器支持实时预览，编辑内容即时渲染",
        "图片上传：支持点击选择和拖拽上传，上传进度实时显示",
        "消息提示：操作结果通过浮动消息提示框反馈给用户",
        "加载状态：数据加载时显示加载动画，提升用户体验",
        "平滑滚动：页面滚动平滑，回到顶部按钮在滚动一定距离后显示",
    ]
    for i in interactions:
        add_para(doc, f"• {i}", font_size=12, space_after=3)
    
    doc.add_page_break()
    
    # ========== 八、安全设计 ==========
    add_heading(doc, "八、安全设计", level=1)
    
    add_heading(doc, "8.1 密码安全", level=2)
    add_para(doc, "用户密码采用bcryptjs库进行哈希加密存储，bcrypt是一种自适应哈希算法，"
             "通过增加计算成本因子来抵抗暴力破解攻击。即使数据库泄露，攻击者也无法还原原始密码。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "8.2 身份认证", level=2)
    add_para(doc, "系统采用JWT（JSON Web Token）进行身份认证。用户登录成功后，服务器生成包含用户ID和用户名的JWT令牌，"
             "设置24小时有效期。客户端在后续请求中通过Authorization头携带令牌，服务器验证令牌的有效性后处理请求。"
             "JWT的无状态特性使得系统易于扩展。", font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "8.3 权限控制", level=2)
    add_para(doc, "系统实现了细粒度的权限控制：", font_size=12)
    perms = [
        "文章操作：只有文章作者才能编辑和删除自己的文章",
        "画廊操作：只有作品作者才能删除自己的作品",
        "个人中心：只有登录用户才能查看自己的个人信息",
        "图片上传：只有登录用户才能上传图片",
    ]
    for p in perms:
        add_para(doc, f"• {p}", font_size=12, space_after=3)
    
    add_heading(doc, "8.4 其他安全措施", level=2)
    others = [
        "文件上传限制：只允许上传JPEG、JPG、PNG、GIF格式的图片文件，大小限制为2MB",
        "SQL注入防护：使用参数化查询（Prepared Statements）防止SQL注入攻击",
        "CORS配置：生产环境下限制跨域请求的来源",
        "错误处理：统一的错误处理中间件，避免敏感信息泄露",
    ]
    for o in others:
        add_para(doc, f"• {o}", font_size=12, space_after=3)
    
    doc.add_page_break()
    
    # ========== 九、部署与运行 ==========
    add_heading(doc, "九、部署与运行", level=1)
    
    add_heading(doc, "9.1 环境要求", level=2)
    env_reqs = [
        "Node.js 18.0 或更高版本",
        "npm 9.0 或更高版本",
        "现代浏览器（Chrome、Edge、Firefox最新版本）",
    ]
    for r in env_reqs:
        add_para(doc, f"• {r}", font_size=12, space_after=3)
    
    add_heading(doc, "9.2 安装与运行", level=2)
    add_para(doc, "第一步：安装依赖", bold=True, font_size=12)
    add_para(doc, "在项目根目录下执行以下命令安装所有依赖包：", font_size=12)
    add_para(doc, "npm install", font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.add_paragraph()
    add_para(doc, "第二步：启动服务器", bold=True, font_size=12)
    add_para(doc, "执行以下命令启动服务器：", font_size=12)
    add_para(doc, "npm start", font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "或使用开发模式（支持热重载）：", font_size=12)
    add_para(doc, "npm run dev", font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.add_paragraph()
    add_para(doc, "第三步：访问网站", bold=True, font_size=12)
    add_para(doc, "打开浏览器，访问 http://localhost:3000 即可进入博客首页。", font_size=12)
    add_para(doc, "默认管理员账号：admin / admin123", font_size=12)
    
    add_heading(doc, "9.3 项目结构", level=2)
    add_para(doc, "my-personal-blog/", bold=True, font_size=11)
    structure = [
        "├── public/              # 前端静态文件",
        "│   ├── index.html       # 主页面",
        "│   ├── gallery.html     # 画廊页面",
        "│   ├── styles.css       # 样式文件",
        "│   ├── script.js        # 前端JavaScript",
        "│   └── uploads/         # 上传文件存储目录",
        "├── server.js            # 服务器入口文件",
        "├── database.js          # 数据库初始化与连接",
        "├── package.json         # 项目配置与依赖",
        "├── .env                 # 环境变量配置",
        "├── blog.db              # SQLite数据库文件",
        "└── README.md            # 项目说明文档",
    ]
    for s in structure:
        add_para(doc, s, font_size=10, space_after=2)
    
    doc.add_page_break()
    
    # ========== 十、总结与展望 ==========
    add_heading(doc, "十、总结与展望", level=1)
    
    add_heading(doc, "10.1 项目总结", level=2)
    add_para(doc, "本系统基于Node.js + Express + SQLite技术栈，成功构建了一个功能完整的个人摄影博客网站。"
             "系统实现了用户管理、文章管理、评论互动、搜索分类、画廊展示等核心功能，"
             "并集成了Markdown编辑器、夜间模式、响应式设计等提升用户体验的特性。", 
             font_size=12, first_line_indent=0.74)
    add_para(doc, "在开发过程中，采用了前后端分离的架构设计，前端使用原生JavaScript实现单页应用效果，"
             "后端提供RESTful API接口。数据库采用SQLite，通过参数化查询确保数据安全。"
             "身份认证采用JWT方案，实现了无状态的用户认证机制。", 
             font_size=12, first_line_indent=0.74)
    add_para(doc, "通过本项目的开发，深入实践了Web全栈开发技术，包括Node.js后端开发、"
             "数据库设计与操作、前端界面设计与交互、API接口设计、用户认证与安全防护等关键技术。", 
             font_size=12, first_line_indent=0.74)
    
    add_heading(doc, "10.2 不足与改进", level=2)
    improvements = [
        "性能优化：可引入Redis缓存机制，减少数据库查询次数，提升响应速度",
        "功能扩展：可增加文章草稿箱、定时发布、多用户角色管理等功能",
        "SEO优化：可引入服务端渲染（SSR）或静态站点生成，提升搜索引擎收录效果",
        "部署方案：可配置Docker容器化部署，简化环境配置和部署流程",
        "测试覆盖：可增加单元测试和集成测试，提高代码质量和稳定性",
        "富文本编辑：可考虑集成更强大的富文本编辑器，如TinyMCE或Quill",
    ]
    for imp in improvements:
        add_para(doc, f"• {imp}", font_size=12, space_after=3)
    
    add_heading(doc, "10.3 展望", level=2)
    add_para(doc, "未来可以进一步扩展系统的功能和应用场景：", font_size=12, first_line_indent=0.74)
    outlooks = [
        "移动端应用：开发微信小程序或移动App，拓展访问渠道",
        "社交分享：集成社交媒体分享功能，扩大内容传播范围",
        "数据分析：增加访问统计和用户行为分析功能，为内容优化提供数据支持",
        "多语言支持：增加国际化功能，支持多语言界面",
        "AI辅助：引入AI写作助手、智能标签推荐等功能，提升创作效率",
    ]
    for o in outlooks:
        add_para(doc, f"• {o}", font_size=12, space_after=3)
    
    # 保存
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "博客网站设计说明书.docx")
    doc.save(output_path)
    print(f"已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 50)
    print("开始生成博客网站设计说明书...")
    print("=" * 50)
    print()
    
    create_design_doc()
    
    print()
    print("=" * 50)
    print("设计说明书已生成完毕！")
    print("=" * 50)
    print()
    print("生成的文件：博客网站设计说明书.docx")
    print("请用Word打开查看和编辑。")
